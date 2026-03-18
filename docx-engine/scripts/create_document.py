#!/usr/bin/env python3
"""
Create document from scratch with content blocks.
Usage: python3 create_document.py config.json output.docx
"""

import sys
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def apply_document_styles(doc, style_config):
    """Apply global document styles."""
    default_font = style_config.get('default_font', 'Calibri')
    default_size = style_config.get('default_size', 11)
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = default_font
    font.size = Pt(default_size)


def create_content_block(doc, block):
    """Create a content block in the document."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    block_type = block.get('type', 'paragraph')
    
    if block_type == 'heading':
        level = block.get('level', 1)
        text = block.get('text', '')
        heading = doc.add_heading(text, level=level)
        
        # Apply custom heading style if provided
        if 'style' in block:
            style = block['style']
            for run in heading.runs:
                if style.get('font'):
                    run.font.name = style['font']
                if style.get('size'):
                    run.font.size = Pt(style['size'])
                if style.get('color'):
                    color = style['color']
                    if color.startswith('#'):
                        color = color[1:]
                    run.font.color.rgb = RGBColor(
                        int(color[0:2], 16),
                        int(color[2:4], 16),
                        int(color[4:6], 16)
                    )
    
    elif block_type == 'paragraph':
        text = block.get('text', '')
        p = doc.add_paragraph(text)
        
        # Apply style
        if 'style' in block:
            style = block['style']
            for run in p.runs:
                if style.get('bold'):
                    run.bold = True
                if style.get('italic'):
                    run.italic = True
                if style.get('font'):
                    run.font.name = style['font']
                if style.get('size'):
                    run.font.size = Pt(style['size'])
        
        # Apply alignment
        if 'alignment' in block:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if block['alignment'] in align_map:
                p.alignment = align_map[block['alignment']]
    
    elif block_type == 'table':
        rows = block.get('rows', [])
        if not rows:
            return
        
        table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
        table.style = block.get('style', 'Table Grid')
        
        # Apply header row formatting if specified
        header_style = block.get('header_style', {})
        
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = str(cell_text)
                    
                    # Apply header formatting to first row
                    if i == 0 and header_style:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if header_style.get('bold'):
                                    run.bold = True
                                if header_style.get('background'):
                                    # Note: python-docx has limited shading support
                                    pass
    
    elif block_type == 'image':
        path = block.get('path', '')
        if path:
            try:
                width = block.get('width')
                height = block.get('height')
                kwargs = {}
                if width:
                    kwargs['width'] = Inches(width)
                if height:
                    kwargs['height'] = Inches(height)
                doc.add_picture(path, **kwargs)
            except Exception as e:
                print(f"Warning: Could not add image {path}: {e}")
    
    elif block_type == 'page_break':
        doc.add_page_break()


def create_document(content_blocks, output_path, style_config=None):
    """Create a DOCX document from content blocks."""
    doc = Document()
    
    # Apply global styles
    if style_config:
        apply_document_styles(doc, style_config)
        
        # Set page size and margins if specified
        if 'page' in style_config:
            page = style_config['page']
            section = doc.sections[0]
            if 'width' in page:
                section.page_width = Inches(page['width'])
            if 'height' in page:
                section.page_height = Inches(page['height'])
            if 'margin_top' in page:
                section.top_margin = Inches(page['margin_top'])
            if 'margin_bottom' in page:
                section.bottom_margin = Inches(page['margin_bottom'])
            if 'margin_left' in page:
                section.left_margin = Inches(page['margin_left'])
            if 'margin_right' in page:
                section.right_margin = Inches(page['margin_right'])
    
    # Add content blocks
    for block in content_blocks:
        create_content_block(doc, block)
    
    # Save document
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Create DOCX from content blocks')
    parser.add_argument('config', help='Path to JSON config file with content blocks')
    parser.add_argument('output', help='Path for output DOCX file')
    
    args = parser.parse_args()
    
    # Load config
    with open(args.config, 'r', encoding='utf-8') as f:
        config = json.load(f)
    
    content_blocks = config.get('content', [])
    style_config = config.get('style', {})
    
    create_document(content_blocks, args.output, style_config)
