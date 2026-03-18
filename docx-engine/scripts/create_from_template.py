#!/usr/bin/env python3
"""
Create document from template with Jinja2 rendering.
Usage: python3 create_from_template.py template.docx data.json output.docx
"""

import sys
import json
from docx import Document
from docx_engine import process_paragraph, process_table, process_headers_footers


def create_from_template(template_path, data, output_path):
    """Create a DOCX from template with Jinja2 data substitution."""
    # Load template
    doc = Document(template_path)
    
    # Process all paragraphs in the document body
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, data)
    
    # Process all tables
    for table in doc.tables:
        process_table(table, data)
    
    # Process headers and footers in all sections
    for section in doc.sections:
        process_headers_footers(section, data)
    
    # Save output
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Create DOCX from template with Jinja2 data')
    parser.add_argument('template', help='Path to template DOCX file')
    parser.add_argument('data', help='Path to JSON data file')
    parser.add_argument('output', help='Path for output DOCX file')
    
    args = parser.parse_args()
    
    # Load data
    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    create_from_template(args.template, data, args.output)
