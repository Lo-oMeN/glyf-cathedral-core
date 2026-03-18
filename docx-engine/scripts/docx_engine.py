#!/usr/bin/env python3
"""
DOCX Template Engine with Jinja2 support.
Core utilities for DOCX templating operations.
"""

import re
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from jinja2 import Environment, BaseLoader


def parse_jinja2_placeholders(text):
    """Extract Jinja2 placeholders from text."""
    if not text:
        return []
    # Match {{ variable }} or {{ variable | filter }}
    pattern = r'\{\{\s*([\w.]+)(?:\s*\|\s*[^}]+)?\s*\}\}'
    return re.findall(pattern, text)


def find_jinja2_loops(text):
    """Find Jinja2 for loops in text."""
    if not text:
        return []
    # Match {% for item in items %}
    pattern = r'\{%\s*for\s+(\w+)\s+in\s+([\w.]+)\s*%\}'
    return re.findall(pattern, text)


def render_jinja2_template(text, data):
    """Render Jinja2 template string with data."""
    if not text:
        return text
    env = Environment(loader=BaseLoader())
    template = env.from_string(text)
    return template.render(**data)


def process_paragraph(paragraph, data):
    """Process a single paragraph, replacing Jinja2 placeholders."""
    full_text = paragraph.text
    if not full_text:
        return
    
    # Check for Jinja2 syntax
    if '{{' in full_text or '{%' in full_text:
        try:
            rendered = render_jinja2_template(full_text, data)
            # Clear and re-add text
            paragraph.clear()
            run = paragraph.add_run(rendered)
            # Preserve basic formatting from first run if exists
            if paragraph.runs:
                run.bold = paragraph.runs[0].bold if paragraph.runs else False
                run.italic = paragraph.runs[0].italic if paragraph.runs else False
        except Exception as e:
            print(f"Warning: Could not render template in paragraph: {e}")


def process_table(table, data):
    """Process table cells for Jinja2 templating."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, data)


def process_headers_footers(section, data):
    """Process headers and footers in a section."""
    # Header
    if section.header:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph, data)
        for table in section.header.tables:
            process_table(table, data)
    
    # Footer
    if section.footer:
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph, data)
        for table in section.footer.tables:
            process_table(table, data)


def apply_style_to_paragraph(paragraph, style_config):
    """Apply styling to a paragraph."""
    font_config = style_config.get('font', {})
    
    for run in paragraph.runs:
        if font_config.get('name'):
            run.font.name = font_config['name']
        if font_config.get('size'):
            run.font.size = Pt(font_config['size'])
        if font_config.get('color'):
            color = font_config['color']
            if isinstance(color, str) and color.startswith('#'):
                color = color[1:]
            if isinstance(color, str) and len(color) == 6:
                run.font.color.rgb = RGBColor(
                    int(color[0:2], 16),
                    int(color[2:4], 16),
                    int(color[4:6], 16)
                )
        if font_config.get('bold') is not None:
            run.bold = font_config['bold']
        if font_config.get('italic') is not None:
            run.italic = font_config['italic']


def set_paragraph_alignment(paragraph, alignment):
    """Set paragraph alignment."""
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    if alignment in align_map:
        paragraph.alignment = align_map[alignment]


# Import additional modules for document creation
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_content_block(doc, block):
    """Create a content block in the document."""
    block_type = block.get('type', 'paragraph')
    
    if block_type == 'heading':
        level = block.get('level', 1)
        text = block.get('text', '')
        doc.add_heading(text, level=level)
    
    elif block_type == 'paragraph':
        text = block.get('text', '')
        p = doc.add_paragraph(text)
        if 'style' in block:
            apply_style_to_paragraph(p, block['style'])
        if 'alignment' in block:
            set_paragraph_alignment(p, block['alignment'])
    
    elif block_type == 'table':
        rows = block.get('rows', [])
        if not rows:
            return
        
        table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
        table.style = block.get('style', 'Table Grid')
        
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = str(cell_text)
        
        if 'alignment' in block:
            table.alignment = WD_TABLE_ALIGNMENT[block.get('alignment', 'LEFT')]
    
    elif block_type == 'image':
        path = block.get('path', '')
        width = block.get('width')
        height = block.get('height')
        
        if path:
            kwargs = {}
            if width:
                kwargs['width'] = Inches(width)
            if height:
                kwargs['height'] = Inches(height)
            doc.add_picture(path, **kwargs)
    
    elif block_type == 'page_break':
        doc.add_page_break()


if __name__ == '__main__':
    print("DOCX Template Engine - Import this module for DOCX operations")
