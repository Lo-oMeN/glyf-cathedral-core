#!/usr/bin/env python3
"""
Extract plain text from DOCX document.
Usage: python3 extract_text.py input.docx > output.txt
       python3 extract_text.py input.docx --output output.txt
"""

import argparse
from docx import Document


def extract_text(doc_path, include_headers=False, include_footers=False):
    """Extract plain text from DOCX document."""
    doc = Document(doc_path)
    
    text_parts = []
    
    # Extract from body paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(' | '.join(row_text))
    
    # Extract from headers if requested
    if include_headers:
        for section in doc.sections:
            if section.header:
                header_text = []
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        header_text.append(paragraph.text)
                if header_text:
                    text_parts.append('[HEADER] ' + ' '.join(header_text))
    
    # Extract from footers if requested
    if include_footers:
        for section in doc.sections:
            if section.footer:
                footer_text = []
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        footer_text.append(paragraph.text)
                if footer_text:
                    text_parts.append('[FOOTER] ' + ' '.join(footer_text))
    
    return '\n\n'.join(text_parts)


def extract_text_simple(doc_path):
    """Simple text extraction - just paragraphs and tables."""
    doc = Document(doc_path)
    
    full_text = []
    
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    return '\n'.join(full_text)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Extract text from DOCX')
    parser.add_argument('input', help='Input DOCX file')
    parser.add_argument('--output', '-o', help='Output text file (default: stdout)')
    parser.add_argument('--headers', action='store_true', help='Include header content')
    parser.add_argument('--footers', action='store_true', help='Include footer content')
    
    args = parser.parse_args()
    
    text = extract_text(args.input, args.headers, args.footers)
    
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Text extracted to: {args.output}")
    else:
        print(text)
