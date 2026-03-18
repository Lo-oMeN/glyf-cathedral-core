#!/usr/bin/env python3
"""
Mail merge - Bulk document generation from template.
Usage: python3 mail_merge.py template.docx data.csv output_dir/
       python3 mail_merge.py template.docx data.json output_dir/
"""

import os
import sys
import csv
import json
from pathlib import Path
from docx import Document
from docx_engine import process_paragraph, process_table, process_headers_footers


def mail_merge(template_path, data_rows, output_dir, filename_field=None):
    """Generate multiple documents from template and data rows."""
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    created_files = []
    
    for i, row in enumerate(data_rows):
        # Generate filename
        if filename_field and filename_field in row:
            filename = f"{row[filename_field]}.docx"
        else:
            filename = f"document_{i+1:04d}.docx"
        
        output_file = output_path / filename
        
        # Load template fresh each iteration
        doc = Document(template_path)
        
        # Process all content
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, row)
        
        for table in doc.tables:
            process_table(table, row)
        
        for section in doc.sections:
            process_headers_footers(section, row)
        
        # Save
        doc.save(output_file)
        created_files.append(output_file)
        print(f"Created: {output_file}")
    
    print(f"\nGenerated {len(created_files)} documents in {output_dir}")
    return created_files


def load_data(data_path):
    """Load data from CSV or JSON file."""
    path = Path(data_path)
    
    if path.suffix.lower() == '.csv':
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            return list(reader)
    
    elif path.suffix.lower() == '.json':
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            # Handle both array and object with 'data' key
            if isinstance(data, list):
                return data
            elif isinstance(data, dict) and 'data' in data:
                return data['data']
            else:
                return [data]  # Single record
    
    else:
        raise ValueError(f"Unsupported file format: {path.suffix}")


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Mail merge - bulk document generation')
    parser.add_argument('template', help='Path to template DOCX file')
    parser.add_argument('data', help='Path to data file (CSV or JSON)')
    parser.add_argument('output_dir', help='Output directory for generated documents')
    parser.add_argument('--filename-field', '-f', help='Field to use for output filenames')
    
    args = parser.parse_args()
    
    # Load data rows
    data_rows = load_data(args.data)
    print(f"Loaded {len(data_rows)} records from {args.data}")
    
    # Generate documents
    mail_merge(args.template, data_rows, args.output_dir, args.filename_field)
