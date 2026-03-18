#!/usr/bin/env python3
"""
Test script for docx-engine skill.
Creates sample templates and tests all operations.
"""

import os
import sys
import json
from pathlib import Path

# Add scripts directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_samples_dir():
    """Get the samples directory path."""
    script_dir = Path(__file__).parent
    return script_dir.parent / 'assets' / 'samples'


def create_invoice_template(output_path):
    """Create a sample invoice template."""
    doc = Document()
    
    # Header
    title = doc.add_heading('INVOICE', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Invoice info
    p = doc.add_paragraph()
    p.add_run('Invoice #: ').bold = True
    p.add_run('{{ invoice_number }}')
    
    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run('{{ date }}')
    
    # Company info
    doc.add_heading('From:', level=2)
    doc.add_paragraph('{{ company.name }}')
    doc.add_paragraph('{{ company.address }}')
    doc.add_paragraph('Email: {{ company.email }}')
    doc.add_paragraph('Phone: {{ company.phone }}')
    
    # Customer info
    doc.add_heading('Bill To:', level=2)
    doc.add_paragraph('{{ customer.name }}')
    doc.add_paragraph('{{ customer.company }}')
    doc.add_paragraph('{{ customer.address }}')
    
    # Items table
    doc.add_heading('Items:', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ['Description', 'Quantity', 'Unit Price', 'Amount']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Template row (this will be duplicated per item)
    row = table.add_row()
    row.cells[0].text = '{{ item.description }}'
    row.cells[1].text = '{{ item.quantity }}'
    row.cells[2].text = '${{ item.unit_price }}'
    row.cells[3].text = '${{ item.quantity * item.unit_price }}'
    
    # Totals
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Subtotal: ').bold = True
    p.add_run('${{ subtotal }}')
    
    p = doc.add_paragraph()
    p.add_run('Tax ({{ tax_rate * 100 }}%): ').bold = True
    p.add_run('${{ tax_amount }}')
    
    p = doc.add_paragraph()
    p.add_run('Total: ').bold = True
    run = p.add_run('${{ total }}')
    run.bold = True
    run.font.size = Pt(14)
    
    # Terms
    doc.add_heading('Payment Terms:', level=2)
    doc.add_paragraph('{{ payment_terms }}')
    
    doc.add_paragraph()
    p = doc.add_paragraph('{{ notes }}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(output_path)
    print(f"Created template: {output_path}")


def create_letter_template(output_path):
    """Create a sample business letter template."""
    doc = Document()
    
    # Date
    doc.add_paragraph('{{ date }}')
    doc.add_paragraph()
    
    # Recipient
    doc.add_paragraph('{{ recipient_name }}')
    doc.add_paragraph('{{ recipient_title }}')
    doc.add_paragraph('{{ recipient_company }}')
    doc.add_paragraph('{{ recipient_address }}')
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph('Dear {{ recipient_name | split(" ") | first }},')
    doc.add_paragraph()
    
    # Body
    doc.add_paragraph('{{ body_paragraph1 }}')
    doc.add_paragraph()
    doc.add_paragraph('{{ body_paragraph2 }}')
    doc.add_paragraph()
    
    # Conditional content
    p = doc.add_paragraph()
    p.add_run('{% if include_offer %}Special Offer: {{ offer_details }}{% endif %}')
    
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph('{{ closing }}')
    doc.add_paragraph()
    doc.add_paragraph('{{ sender_name }}')
    doc.add_paragraph('{{ sender_title }}')
    
    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page {{ page_number }} | {{ company_name }}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(output_path)
    print(f"Created template: {output_path}")


def test_create_from_template():
    """Test template rendering."""
    print("\n=== Test: create_from_template ===")
    
    from docx_engine import process_paragraph, process_table, process_headers_footers
    
    samples_dir = get_samples_dir()
    
    # Load data
    with open(samples_dir / 'sample_data.json', 'r') as f:
        data = json.load(f)
    
    # Process template
    doc = Document(samples_dir / 'invoice_template.docx')
    
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, data)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, data)
    
    output_path = samples_dir / 'test_invoice_output.docx'
    doc.save(output_path)
    print(f"Output saved: {output_path}")


def test_create_document():
    """Test creating document from scratch."""
    print("\n=== Test: create_document ===")
    
    from create_document import create_document
    
    samples_dir = get_samples_dir()
    
    config = {
        "style": {
            "default_font": "Arial",
            "default_size": 11,
            "page": {
                "margin_top": 1,
                "margin_bottom": 1,
                "margin_left": 1,
                "margin_right": 1
            }
        },
        "content": [
            {
                "type": "heading",
                "level": 1,
                "text": "Test Document"
            },
            {
                "type": "paragraph",
                "text": "This is a test document created from content blocks.",
                "alignment": "left"
            },
            {
                "type": "heading",
                "level": 2,
                "text": "Data Table"
            },
            {
                "type": "table",
                "rows": [
                    ["Name", "Value", "Status"],
                    ["Item 1", "100", "Active"],
                    ["Item 2", "200", "Inactive"]
                ],
                "style": "Table Grid",
                "header_style": {"bold": True}
            },
            {
                "type": "paragraph",
                "text": "End of document."
            }
        ]
    }
    
    output_path = samples_dir / 'test_created.docx'
    create_document(config['content'], str(output_path), config['style'])


def test_mail_merge():
    """Test mail merge functionality."""
    print("\n=== Test: mail_merge ===")
    
    from mail_merge import mail_merge, load_data
    
    samples_dir = get_samples_dir()
    
    # Load data
    data_rows = load_data(samples_dir / 'mail_merge_data.json')
    
    # Create a simple template for mail merge
    doc = Document()
    doc.add_heading('Certificate', level=1)
    doc.add_paragraph('This certifies that {{ name }}')
    doc.add_paragraph('from {{ company }}')
    doc.add_paragraph('has achieved {{ status }} status.')
    
    template_path = samples_dir / 'cert_template.docx'
    doc.save(template_path)
    
    output_dir = samples_dir / 'merged'
    output_dir.mkdir(exist_ok=True)
    
    mail_merge(str(template_path), data_rows, str(output_dir), filename_field='name')


def test_extract_text():
    """Test text extraction."""
    print("\n=== Test: extract_text ===")
    
    from extract_text import extract_text
    
    samples_dir = get_samples_dir()
    
    text = extract_text(samples_dir / 'test_invoice_output.docx')
    print("Extracted text preview (first 500 chars):")
    print(text[:500])
    print("...")


def run_all_tests():
    """Run all tests."""
    print("DOCX Engine - Test Suite")
    print("=" * 50)
    
    # Get the correct base path
    samples_dir = get_samples_dir()
    samples_dir.mkdir(parents=True, exist_ok=True)
    
    # Create templates first
    create_invoice_template(samples_dir / 'invoice_template.docx')
    create_letter_template(samples_dir / 'letter_template.docx')
    
    # Run tests
    test_create_from_template()
    test_create_document()
    test_mail_merge()
    test_extract_text()
    
    print("\n" + "=" * 50)
    print("All tests completed!")


if __name__ == '__main__':
    run_all_tests()
