#!/usr/bin/env python3
"""
Report Builder - Automated report generation from data to PDF/Word/HTML with charts.
"""

import json
import csv
import os
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# Report generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.pdfgen import canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Non-interactive backend
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    import plotly.graph_objects as go
    import plotly.io as pio
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False


class Report:
    """Report object that holds content and configuration."""
    
    def __init__(self, title: str = "Report", template: str = "default", output_format: str = "pdf"):
        self.title = title
        self.template = template
        self.output_format = output_format.lower()
        self.sections = []
        self.charts = []
        self.tables = []
        self.metadata = {
            "created_at": datetime.now().isoformat(),
            "author": "Report Builder"
        }
    
    def add_section(self, title: str, content: str):
        """Add a text section to the report."""
        self.sections.append({
            "type": "section",
            "title": title,
            "content": content
        })
    
    def add_chart(self, chart_type: str, data: Dict, title: str = ""):
        """Add a chart to the report."""
        self.charts.append({
            "type": chart_type,
            "data": data,
            "title": title
        })
    
    def add_table(self, data: List[List], headers: Optional[List[str]] = None):
        """Add a table to the report."""
        self.tables.append({
            "data": data,
            "headers": headers
        })


def create_report(data: Union[Dict, List, str], template: str = "default", output_format: str = "pdf") -> Report:
    """
    Create a new report from data.
    
    Args:
        data: Input data (dict, list, or path to JSON/CSV file)
        template: Report template (default, executive, technical, financial)
        output_format: Output format (pdf, docx, html)
    
    Returns:
        Report object
    """
    # Load data if path provided
    if isinstance(data, str) and os.path.isfile(data):
        data = load_data_source(data)
    
    # Create report with template
    report = Report(
        title=data.get("title", "Report") if isinstance(data, dict) else "Report",
        template=template,
        output_format=output_format
    )
    
    # Apply template defaults
    if template == "executive":
        apply_executive_template(report, data)
    elif template == "technical":
        apply_technical_template(report, data)
    elif template == "financial":
        apply_financial_template(report, data)
    else:
        # Default template - just add content sections
        if isinstance(data, dict):
            for key, value in data.items():
                if key != "title" and isinstance(value, str):
                    report.add_section(key.replace("_", " ").title(), value)
    
    return report


def add_chart(report: Report, chart_type: str, data: Dict, title: str = ""):
    """
    Add a chart to the report.
    
    Args:
        report: Report object
        chart_type: Type of chart (bar, line, pie, scatter)
        data: Chart data (dict with 'x' and 'y' keys, or 'labels' and 'values' for pie)
        title: Chart title
    """
    report.add_chart(chart_type, data, title)


def add_table(report: Report, data: List[List], headers: Optional[List[str]] = None):
    """
    Add a table to the report.
    
    Args:
        report: Report object
        data: Table data (list of lists)
        headers: Optional column headers
    """
    report.add_table(data, headers)


def add_section(report: Report, title: str, content: str):
    """
    Add a section to the report.
    
    Args:
        report: Report object
        title: Section title
        content: Section content
    """
    report.add_section(title, content)


def export(report: Report, output_path: str, format: Optional[str] = None) -> str:
    """
    Export the report to a file.
    
    Args:
        report: Report object
        output_path: Output file path
        format: Output format (pdf, docx, html). Defaults to report's format.
    
    Returns:
        Path to the exported file
    """
    fmt = format or report.output_format
    fmt = fmt.lower()
    
    if fmt == "docx":
        return export_to_docx(report, output_path)
    elif fmt == "html":
        return export_to_html(report, output_path)
    else:  # Default to PDF
        return export_to_pdf(report, output_path)


def load_data_source(source: str) -> Union[Dict, List]:
    """
    Load data from a file (JSON or CSV).
    
    Args:
        source: Path to data file
    
    Returns:
        Loaded data (dict or list)
    """
    ext = Path(source).suffix.lower()
    
    if ext == ".json":
        with open(source, 'r') as f:
            return json.load(f)
    elif ext == ".csv":
        data = []
        with open(source, 'r') as f:
            reader = csv.DictReader(f)
            for row in reader:
                data.append(row)
        return data
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def generate_chart_image(chart_data: Dict, output_path: str) -> str:
    """Generate a chart image using matplotlib."""
    if not HAS_MATPLOTLIB:
        raise ImportError("matplotlib is required for chart generation")
    
    chart_type = chart_data.get("type", "bar")
    data = chart_data.get("data", {})
    title = chart_data.get("title", "")
    
    plt.figure(figsize=(8, 5))
    
    if chart_type == "bar":
        x = data.get("x", [])
        y = data.get("y", [])
        plt.bar(x, y, color='steelblue')
        plt.xlabel(data.get("xlabel", ""))
        plt.ylabel(data.get("ylabel", ""))
    
    elif chart_type == "line":
        x = data.get("x", [])
        y = data.get("y", [])
        plt.plot(x, y, marker='o', color='steelblue')
        plt.xlabel(data.get("xlabel", ""))
        plt.ylabel(data.get("ylabel", ""))
    
    elif chart_type == "pie":
        labels = data.get("labels", [])
        values = data.get("values", [])
        plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
    
    elif chart_type == "scatter":
        x = data.get("x", [])
        y = data.get("y", [])
        plt.scatter(x, y, color='steelblue')
        plt.xlabel(data.get("xlabel", ""))
        plt.ylabel(data.get("ylabel", ""))
    
    plt.title(title)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    
    return output_path


def apply_executive_template(report: Report, data: Union[Dict, List]):
    """Apply executive summary template."""
    if isinstance(data, dict):
        # Executive summary section
        summary = data.get("executive_summary", data.get("summary", ""))
        if summary:
            report.add_section("Executive Summary", summary)
        
        # Key findings
        findings = data.get("key_findings", [])
        if findings:
            content = "\n".join(f"• {finding}" for finding in findings)
            report.add_section("Key Findings", content)
        
        # Recommendations
        recommendations = data.get("recommendations", [])
        if recommendations:
            content = "\n".join(f"• {rec}" for rec in recommendations)
            report.add_section("Recommendations", content)


def apply_technical_template(report: Report, data: Union[Dict, List]):
    """Apply technical report template."""
    if isinstance(data, dict):
        # Overview
        overview = data.get("overview", data.get("introduction", ""))
        if overview:
            report.add_section("Overview", overview)
        
        # Methodology
        methodology = data.get("methodology", data.get("approach", ""))
        if methodology:
            report.add_section("Methodology", methodology)
        
        # Results
        results = data.get("results", data.get("findings", ""))
        if results:
            report.add_section("Results", results)
        
        # Conclusion
        conclusion = data.get("conclusion", data.get("conclusions", ""))
        if conclusion:
            report.add_section("Conclusion", conclusion)


def apply_financial_template(report: Report, data: Union[Dict, List]):
    """Apply financial report template."""
    if isinstance(data, dict):
        # Period
        period = data.get("period", data.get("reporting_period", ""))
        if period:
            report.add_section("Reporting Period", period)
        
        # Financial highlights
        highlights = data.get("highlights", data.get("financial_highlights", []))
        if highlights:
            content = "\n".join(f"• {highlight}" for highlight in highlights)
            report.add_section("Financial Highlights", content)
        
        # Revenue summary
        revenue = data.get("revenue", {})
        if revenue:
            content = f"""
Total Revenue: {revenue.get('total', 'N/A')}
Growth: {revenue.get('growth', 'N/A')}
Breakdown: {', '.join(revenue.get('breakdown', []))}
"""
            report.add_section("Revenue Summary", content)


def export_to_pdf(report: Report, output_path: str) -> str:
    """Export report to PDF using reportlab."""
    if not HAS_REPORTLAB:
        raise ImportError("reportlab is required for PDF export")
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )
    
    styles = getSampleStyleSheet()
    story = []
    temp_images = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=WD_ALIGN_PARAGRAPH.CENTER if hasattr(WD_ALIGN_PARAGRAPH, 'CENTER') else 1
    )
    story.append(Paragraph(report.title, title_style))
    story.append(Spacer(1, 12))
    
    # Metadata
    meta_text = f"Generated: {report.metadata['created_at'][:10]}"
    story.append(Paragraph(meta_text, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Sections
    for section in report.sections:
        story.append(Paragraph(section['title'], styles['Heading2']))
        story.append(Spacer(1, 6))
        
        # Handle multi-line content
        for line in section['content'].split('\n'):
            if line.strip().startswith('•'):
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{line}", styles['Normal']))
            else:
                story.append(Paragraph(line, styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Tables
    for table_data in report.tables:
        story.append(Spacer(1, 12))
        
        headers = table_data.get('headers', [])
        data = table_data.get('data', [])
        
        if headers:
            table_content = [headers] + data
        else:
            table_content = data
        
        if table_content:
            table = Table(table_content)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
            story.append(Spacer(1, 12))
    
    # Charts
    for chart in report.charts:
        story.append(Spacer(1, 12))
        
        # Generate chart image
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            chart_path = tmp.name
        
        try:
            generate_chart_image(chart, chart_path)
            temp_images.append(chart_path)
            
            if chart.get('title'):
                story.append(Paragraph(chart['title'], styles['Heading3']))
            
            img = Image(chart_path, width=400, height=250)
            story.append(img)
            story.append(Spacer(1, 12))
        except Exception as e:
            story.append(Paragraph(f"[Chart Error: {e}]", styles['Normal']))
    
    # Build PDF
    doc.build(story)
    
    # Cleanup temp images
    for img_path in temp_images:
        try:
            os.unlink(img_path)
        except:
            pass
    
    return output_path


def export_to_docx(report: Report, output_path: str) -> str:
    """Export report to Word document."""
    if not HAS_DOCX:
        raise ImportError("python-docx is required for DOCX export")
    
    doc = Document()
    temp_images = []
    
    # Title
    title = doc.add_heading(report.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {report.metadata['created_at'][:10]}")
    doc.add_paragraph()
    
    # Sections
    for section in report.sections:
        doc.add_heading(section['title'], level=1)
        
        for line in section['content'].split('\n'):
            if line.strip():
                doc.add_paragraph(line)
    
    # Tables
    for table_data in report.tables:
        headers = table_data.get('headers', [])
        data = table_data.get('data', [])
        
        if headers:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = str(header)
            
            # Data rows
            for row_data in data:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = str(cell_data)
        else:
            if data:
                table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 1)
                table.style = 'Light Grid Accent 1'
                
                for i, row_data in enumerate(data):
                    for j, cell_data in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = str(cell_data)
        
        doc.add_paragraph()
    
    # Charts
    for chart in report.charts:
        # Generate chart image
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            chart_path = tmp.name
        
        try:
            generate_chart_image(chart, chart_path)
            temp_images.append(chart_path)
            
            if chart.get('title'):
                doc.add_heading(chart['title'], level=2)
            
            doc.add_picture(chart_path, width=Inches(5.5))
            doc.add_paragraph()
        except Exception as e:
            doc.add_paragraph(f"[Chart Error: {e}]")
    
    # Save document
    doc.save(output_path)
    
    # Cleanup temp images
    for img_path in temp_images:
        try:
            os.unlink(img_path)
        except:
            pass
    
    return output_path


def export_to_html(report: Report, output_path: str) -> str:
    """Export report to HTML."""
    temp_images = []
    
    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        f"<title>{report.title}</title>",
        "<style>",
        "body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }",
        "h1 { color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }",
        "h2 { color: #555; margin-top: 30px; }",
        "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
        "th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }",
        "th { background-color: #4CAF50; color: white; }",
        "tr:nth-child(even) { background-color: #f2f2f2; }",
        ".chart { margin: 20px 0; text-align: center; }",
        ".chart img { max-width: 100%; height: auto; }",
        "</style>",
        "</head>",
        "<body>",
        f"<h1>{report.title}</h1>",
        f"<p><em>Generated: {report.metadata['created_at'][:10]}</em></p>",
    ]
    
    # Sections
    for section in report.sections:
        html_parts.append(f"<h2>{section['title']}</h2>")
        content = section['content'].replace('\n', '<br>')
        html_parts.append(f"<p>{content}</p>")
    
    # Tables
    for table_data in report.tables:
        html_parts.append("<table>")
        
        headers = table_data.get('headers', [])
        if headers:
            html_parts.append("<tr>" + "".join(f"<th>{h}</th>" for h in headers) + "</tr>")
        
        for row in table_data.get('data', []):
            html_parts.append("<tr>" + "".join(f"<td>{cell}</td>" for cell in row) + "</tr>")
        
        html_parts.append("</table>")
    
    # Charts
    for chart in report.charts:
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            chart_path = tmp.name
        
        try:
            generate_chart_image(chart, chart_path)
            temp_images.append(chart_path)
            
            # Convert to base64 for inline embedding
            import base64
            with open(chart_path, 'rb') as f:
                img_data = base64.b64encode(f.read()).decode()
            
            html_parts.append("<div class='chart'>")
            if chart.get('title'):
                html_parts.append(f"<h3>{chart['title']}</h3>")
            html_parts.append(f"<img src='data:image/png;base64,{img_data}' />")
            html_parts.append("</div>")
        except Exception as e:
            html_parts.append(f"<p>[Chart Error: {e}]</p>")
    
    html_parts.extend(["</body>", "</html>"])
    
    # Write HTML
    with open(output_path, 'w') as f:
        f.write('\n'.join(html_parts))
    
    # Cleanup temp images
    for img_path in temp_images:
        try:
            os.unlink(img_path)
        except:
            pass
    
    return output_path


# CLI interface
def main():
    """CLI entry point for testing."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Report Builder')
    parser.add_argument('--data', '-d', help='Path to data file (JSON/CSV)')
    parser.add_argument('--template', '-t', default='default', 
                        choices=['default', 'executive', 'technical', 'financial'],
                        help='Report template')
    parser.add_argument('--format', '-f', default='pdf',
                        choices=['pdf', 'docx', 'html'],
                        help='Output format')
    parser.add_argument('--output', '-o', required=True, help='Output file path')
    
    args = parser.parse_args()
    
    # Create report
    data = args.data or {}
    report = create_report(data, template=args.template, output_format=args.format)
    
    # Export
    result_path = export(report, args.output, args.format)
    print(f"Report exported to: {result_path}")


if __name__ == "__main__":
    main()
